from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/enterprise_service_knowledge_base_longform.docx")


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 116, 181)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 116, 181)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        set_east_asia_font(style, "PingFang SC")
        if style_name == "Normal":
            style.font.name = "Arial"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.25


CHAPTERS = [
    (
        "第一章 企业服务定位与沟通原则",
        "企业服务定位",
        "品牌承诺、服务语气和解释边界",
        ["企业能提供什么服务", "客服身份说明", "处理依据解释"],
    ),
    (
        "第二章 用户身份、称呼与隐私保护",
        "身份与隐私",
        "用户称呼、账号标识和敏感信息保护",
        ["用户更改称呼", "用户提供会员账号", "用户要求使用历史地址"],
    ),
    (
        "第三章 商品信息、价格解释与购买咨询",
        "商品与价格",
        "价格来源、商品规格和购买意向确认",
        ["商品价格咨询", "商品规格对比", "购买前确认"],
    ),
    (
        "第四章 订单创建、支付确认与取消处理",
        "订单处理",
        "下单确认、支付状态和取消窗口",
        ["创建订单", "支付核对", "取消刚创建的订单"],
    ),
    (
        "第五章 物流履约、配送承诺与改派协同",
        "物流履约",
        "发货进度、配送承诺和地址改派",
        ["订单未发货", "配送晚到", "用户要求改地址"],
    ),
    (
        "第六章 退款、退货与换货处理",
        "售后处理",
        "退款条件、退货回收和换货履约",
        ["申请退款", "申请退货", "换颜色或换规格"],
    ),
    (
        "第七章 会员等级、权益发放与补偿",
        "会员权益",
        "等级资格、权益发放和补偿边界",
        ["会员券未到账", "积分缺失", "活动礼品延迟"],
    ),
    (
        "第八章 投诉、风险与人工介入",
        "投诉风险",
        "升级路径、证据保全和响应时限",
        ["用户投诉", "要求人工", "涉及隐私或扣款风险"],
    ),
    (
        "第九章 企业文化、品牌语气与服务红线",
        "企业文化",
        "服务语气、禁用表达和承诺边界",
        ["用户质疑态度", "要求保证结果", "要求内部规则"],
    ),
    (
        "第十章 内部运营排查与数据口径",
        "运营排查",
        "指标口径、工单归因和复盘字段",
        ["差评上升", "责任归因", "处理耗时复盘"],
    ),
    (
        "第十一章 知识维护、工具协作与流程治理",
        "流程治理",
        "知识维护、工具调用和流程更新",
        ["文档更新", "接口能力变化", "流程发布回滚"],
    ),
    (
        "第十二章 服务质量复盘与持续改进",
        "质量改进",
        "质量检查、问题复盘和改进闭环",
        ["处理结果复盘", "服务标准调整", "跨团队协作改进"],
    ),
]


def chapter_paragraphs(chapter: str, domain: str, concern: str, examples: list[str]) -> list[str]:
    example_text = "、".join(examples)
    return [
        (
            f"{chapter}用于统一{domain}相关事项的处理方式。服务人员接到{example_text}等问题时，"
            f"应先确认用户真实诉求、当前已知事实、可执行动作和需要补充的信息。若用户表达中同时包含多个事项，"
            f"应区分主诉求与后续诉求，先处理时效更强、风险更高或用户明确要求优先处理的部分。"
        ),
        (
            f"{domain}场景下的回复需要同时满足准确、清楚和可执行三项要求。准确是指结论必须来自已确认事实、"
            f"已发布规则或可追溯记录；清楚是指用户能理解当前状态和下一步；可执行是指回复后能形成追问、查询、"
            f"创建、取消、补偿、转人工或结束等明确动作。"
        ),
        (
            f"涉及{concern}时，服务人员应避免过度承诺。能够立即办理的事项，应说明已办理结果和后续影响；"
            f"需要等待外部处理的事项，应说明预计观察点、责任团队和用户可以采取的下一步；需要补充信息的事项，"
            f"应一次性列出关键缺口，避免连续多轮只追问一个字段。"
        ),
        (
            f"如果同一问题存在多个规则来源，应优先采用发布时间较新、适用范围较窄、与用户条件更匹配的内容。"
            f"当规则之间存在明显冲突时，不应自行扩大解释，而应保留冲突点并转给相应负责人确认。"
            f"对用户侧表达，应使用结论先行的语言，内部原因只在必要时简要说明。"
        ),
        (
            f"{domain}事项处理完成后，应留下能够复盘的摘要，包括用户诉求、关键事实、采取动作、未解决风险和后续责任。"
            f"若用户继续提出新事项，应在原事项完成状态明确后再进入下一事项，避免把不同事项的字段、工具结果或处理结论混在一起。"
        ),
    ]


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("面壁智能客户服务知识手册")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("企业介绍、服务原则、订单履约、会员权益与售后处理参考")
    subtitle_run.font.name = "Arial"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_paragraph(
        doc,
        "版本：2026-06-16。本文档面向客户服务、运营管理、产品支持和服务质量团队，"
        "汇总企业客户服务中的常见规则、处理原则、沟通边界和跨团队协作方式。"
        "文档以章节和段落组织，适合长期维护、定期复盘和按业务主题扩展。",
    )
    add_heading(doc, "文档说明", 1)
    for item in [
        "本手册适用于客户服务、运营复盘、流程治理和知识库维护。",
        "处理用户问题时，应优先区分事实核对、政策解释、执行动作和风险升级。",
        "涉及敏感信息、补偿承诺或人工介入时，应保留必要依据并避免过度承诺。",
        "当用户一次提出多个诉求时，应先明确优先级，再按事项形成独立处理结论。",
    ]:
        add_bullet(doc, item)

    for chapter_index, (chapter, domain, concern, examples) in enumerate(CHAPTERS, start=1):
        add_heading(doc, chapter, 1)
        add_paragraph(
            doc,
            f"本章围绕{domain}展开，关注{concern}。相关内容适用于日常客服对话、内部工单流转、"
            "运营排查和服务质量复盘。服务人员应根据用户当前诉求和已确认事实选择合适处理方式。",
        )
        for section_index in range(1, 5):
            add_heading(doc, f"{chapter_index}.{section_index} {domain}处理原则 {section_index}", 2)
            for paragraph in chapter_paragraphs(chapter, domain, concern, examples):
                add_paragraph(
                    doc,
                    paragraph
                    + f"在第 {chapter_index} 章第 {section_index} 节的具体执行中，还应结合当前业务状态、"
                    f"用户历史偏好、可用工具返回、知识库记录和人工审核要求，形成清晰的处理链路。"
                )
            add_heading(doc, f"{chapter_index}.{section_index}.1 常见场景说明", 3)
            for item in [
                f"当用户直接描述{examples[0]}时，应先确认是否已有可执行信息，再决定查询、办理或解释。",
                f"当用户同时涉及{examples[1]}和其他事项时，应拆清事项边界，避免把不同处理结果混用。",
                f"当用户表达{examples[2]}时，应说明当前可做动作、限制条件和下一步责任归属。",
            ]:
                add_bullet(doc, item)

    add_heading(doc, "跨章节协作原则", 1)
    cross_domain_topics = [
        (
            "会员权益与订单售后经常同时出现。例如用户反馈会员券未到账又要求取消订单时，"
            "应先确认订单状态，再核对权益资格和发放记录，最后说明取消订单对权益的影响。"
        ),
        (
            "物流履约与补偿承诺也经常交叉。用户认为配送晚到时，需要确认承诺来源、发货状态、"
            "实际签收时间和补偿规则，不应只凭用户口述立即承诺补偿。"
        ),
        (
            "投诉升级与人工介入需要保留完整上下文。转交前应整理用户诉求、已核对字段、已尝试动作、"
            "未解决风险和建议处理方向，减少用户重复描述。"
        ),
        (
            "服务质量复盘应同时查看对话过程和业务结果。单次负面反馈只能说明用户在该轮体验不满意，"
            "还需要结合回复是否准确、流程是否完整、工具是否可用和最终问题是否解决。"
        ),
    ]
    for repeat in range(16):
        for topic in cross_domain_topics:
            add_paragraph(
                doc,
                f"跨章节协作说明 {repeat + 1}：{topic}"
                "处理跨域事项时，应保持事项独立、证据清楚、结论明确，并在用户可理解的范围内说明下一步。"
            )

    add_heading(doc, "服务知识维护原则", 1)
    maintenance_topics = [
        "知识内容应按业务负责人、更新时间、适用范围和例外条件进行维护。过期内容应及时下线，但历史版本仍需保留用于审计。",
        "工具能力发生变化时，应同步更新相关流程说明、可执行动作和异常处理方式，避免前端说明与后端能力不一致。",
        "不同智能体可以拥有不同可见范围。分支中的知识或技能改动不应影响整体版本，除非经过负责人确认并推送到整体。",
        "新规则发布前应检查是否影响现有售后、会员、物流和订单流程，尤其要关注金额、承诺、隐私和人工介入边界。",
    ]
    for repeat in range(12):
        for topic in maintenance_topics:
            add_paragraph(
                doc,
                f"维护说明 {repeat + 1}：{topic}"
                "维护动作完成后，应记录变更原因、影响范围和回滚方式，便于后续复盘。"
            )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    print(f"written={OUTPUT}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"cjk_chars={cjk_count}")


if __name__ == "__main__":
    main()
